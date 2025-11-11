from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_statuts_template() -> None:
    doc = Document()

    sections = doc.sections
    for section in sections:
      section.top_margin = Inches(1)
      section.bottom_margin = Inches(1)
      section.left_margin = Inches(1)
      section.right_margin = Inches(1)

    def add_title(text: str, level: int = 1):
        if level == 0:
            heading = doc.add_heading(text, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return heading
        return doc.add_heading(text, level=level)

    def add_para(text: str, bold: bool = False):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(11)
        return paragraph

    add_title("STATUTS", 0)
    add_title("Société par Actions Simplifiée", 0)
    doc.add_paragraph()

    add_para("Les soussignés :", bold=True)
    add_para("{{associe_civilite}} {{associe_prenom}} {{associe_nom}}")
    add_para("Né(e) le {{associe_date_naissance}} à {{associe_lieu_naissance}}")
    add_para("Demeurant {{associe_adresse_complete}}")
    doc.add_paragraph()
    add_para(
        "Ont établi ainsi qu'il suit les statuts d'une société par actions simplifiée qu'ils déclarent constituer entre eux."
    )
    doc.add_paragraph()

    add_title("TITRE I : FORME - DÉNOMINATION - SIÈGE - OBJET - DURÉE", 1)

    add_title("Article 1 - FORME", 2)
    add_para(
        "Il est formé une {{forme_juridique_complete}} ({{forme_juridique_sigle}}) régie par les dispositions législatives et réglementaires en vigueur, et par les présents statuts."
    )

    add_title("Article 2 - DÉNOMINATION", 2)
    add_para("La société a pour dénomination sociale : {{denomination}}")

    add_title("Article 3 - SIÈGE SOCIAL", 2)
    add_para("Le siège social est fixé à : {{adresse_siege_complete}}")
    add_para(
        "Il pourra être transféré en tout autre endroit par décision du président, sous réserve de ratification par l'assemblée générale."
    )

    add_title("Article 4 - OBJET", 2)
    add_para("La société a pour objet :")
    add_para("{{objet_social}}")
    add_para(
        "Et généralement, toutes opérations industrielles, commerciales, financières, mobilières ou immobilières se rattachant directement ou indirectement à cet objet ou susceptibles d'en faciliter la réalisation."
    )

    add_title("Article 5 - DURÉE", 2)
    add_para(
        "La durée de la société est fixée à {{duree_societe}} années à compter de son immatriculation au Registre du Commerce et des Sociétés, sauf dissolution anticipée ou prorogation."
    )

    add_title("TITRE II : APPORTS - CAPITAL SOCIAL", 1)

    add_title("Article 6 - APPORTS", 2)
    add_para("Les associés font apport à la société de :")
    add_para("- Apports en numéraire : {{capital_social_formate}} euros")

    add_title("Article 7 - CAPITAL SOCIAL", 2)
    add_para("Le capital social est fixé à la somme de {{capital_social_formate}} euros.")
    add_para(
        "Il est divisé en {{nombre_actions}} actions de {{valeur_nominale_formate}} euros chacune, entièrement souscrites et libérées à hauteur de {{montant_libere_formate}} euros."
    )

    add_title("Article 8 - MODIFICATION DU CAPITAL", 2)
    add_para("Le capital social peut être augmenté, réduit ou amorti dans les conditions prévues par la loi.")

    add_title("Article 9 - FORME DES ACTIONS", 2)
    add_para("Les actions sont nominatives. Elles donnent lieu à une inscription en compte.")

    add_title("TITRE III : ADMINISTRATION - DIRECTION", 1)

    add_title("Article 10 - PRÉSIDENCE", 2)
    add_para("La société est représentée par un président désigné parmi les associés ou en dehors d'eux.")
    add_para("Le premier président est : {{associe_civilite}} {{associe_prenom}} {{associe_nom}}")
    add_para("Durée du mandat : Le président est nommé pour une durée illimitée.")

    add_title("Article 11 - POUVOIRS DU PRÉSIDENT", 2)
    add_para(
        "Le président est investi des pouvoirs les plus étendus pour agir en toute circonstance au nom de la société, dans la limite de l'objet social."
    )

    add_title("Article 12 - RÉMUNÉRATION", 2)
    add_para("L'assemblée générale peut allouer au président une rémunération fixe ou proportionnelle.")

    add_title("TITRE IV : DÉCISIONS COLLECTIVES", 1)

    add_title("Article 13 - ASSEMBLÉES GÉNÉRALES", 2)
    add_para(
        "Les associés sont réunis en assemblée générale aussi souvent que l'intérêt de la société l'exige et au moins une fois par an."
    )

    add_title("Article 14 - CONVOCATION", 2)
    add_para("Les associés sont convoqués par le président par tous moyens (lettre simple, email, etc.).")

    add_title("Article 15 - QUORUM ET MAJORITÉ", 2)
    add_para("Chaque action donne droit à une voix.")
    add_para("Les décisions sont prises à la majorité des voix exprimées, sauf dispositions légales contraires.")

    add_title("Article 16 - PROCÈS-VERBAUX", 2)
    add_para("Les décisions sont constatées par des procès-verbaux signés par le président et conservés au siège social.")

    add_title("TITRE V : EXERCICE SOCIAL - COMPTES SOCIAUX", 1)

    add_title("Article 17 - EXERCICE SOCIAL", 2)
    add_para("L'exercice social commence le 1er janvier et se termine le 31 décembre de chaque année.")
    add_para(
        "Par exception, le premier exercice commencera à la date d'immatriculation et se terminera le {{premier_exercice_fin}}."
    )

    add_title("Article 18 - COMPTES ANNUELS", 2)
    add_para("Le président établit les comptes annuels conformément à la loi.")

    add_title("Article 19 - AFFECTATION DU RÉSULTAT", 2)
    add_para("Le bénéfice distribuable est réparti entre les associés proportionnellement au nombre d'actions détenues.")

    add_title("TITRE VI : DISSOLUTION - LIQUIDATION", 1)

    add_title("Article 20 - DISSOLUTION", 2)
    add_para(
        "La société prend fin par l'arrivée du terme, par décision de l'assemblée générale ou pour toute autre cause prévue par la loi."
    )

    add_title("Article 21 - LIQUIDATION", 2)
    add_para("En cas de dissolution, un ou plusieurs liquidateurs sont désignés par l'assemblée générale.")

    doc.add_page_break()
    add_para("Fait à {{adresse_siege_complete}}", bold=True)
    add_para("Le {{date_signature}}", bold=True)
    doc.add_paragraph()
    add_para("Signature du ou des associés :")
    doc.add_paragraph()
    doc.add_paragraph()
    add_para("{{associe_prenom}} {{associe_nom}}")

    doc.save("templates/template-statuts-auto.docx")

    print("✅ Template créé : templates/template-statuts-auto.docx")
    print("📋 Placeholders utilisés:")
    placeholders = [
        "associe_civilite",
        "associe_prenom",
        "associe_nom",
        "associe_date_naissance",
        "associe_lieu_naissance",
        "associe_adresse_complete",
        "forme_juridique_complete",
        "forme_juridique_sigle",
        "denomination",
        "adresse_siege_complete",
        "objet_social",
        "duree_societe",
        "capital_social_formate",
        "nombre_actions",
        "valeur_nominale_formate",
        "montant_libere_formate",
        "premier_exercice_fin",
        "date_signature",
    ]
    for placeholder in placeholders:
        print(f"  - {{{{{placeholder}}}}}")

if __name__ == "__main__":
    create_statuts_template()
