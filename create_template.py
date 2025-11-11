"""
Script pour transformer le document Word en template avec placeholders
"""
from docx import Document
import os

# Chemins des fichiers
input_file = os.path.expanduser("~/Desktop/Statuts SAHEL TRANSPORT.docx")
output_file = "./templates/template-statuts.docx"

# Vérifier que le fichier source existe
if not os.path.exists(input_file):
    print(f"❌ Fichier source introuvable : {input_file}")
    print("📋 Télécharge d'abord le fichier sur ton Bureau !")
    exit(1)

# Charger le document
print("📄 Chargement du document Word...")
doc = Document(input_file)

# Dictionnaire des remplacements
replacements = {
    # ASSOCIÉ
    "Monsieur DIAOU Mamadou": "{{associe_civilite}} {{associe_prenom}} {{associe_nom}}",
    "DIAOU Mamadou": "{{associe_prenom}} {{associe_nom}}",
    "Né (date) à Ville (FRANCE)": "Né le {{associe_date_naissance}} à {{associe_lieu_naissance}} ({{associe_pays_naissance}})",
    "Demeurant au (ADRESSE)": "Demeurant {{associe_adresse_complete}}",
    
    # SOCIÉTÉ
    "SAHEL TRANSPORT": "{{denomination}}",
    "Exportation de marchandises": "{{objet_social_ligne1}}",
    "Et plus généralement toutes opérations commerciales, financières, mobilières, transport léger -3.5t, transport lourd +3.5t, commissionnaire de transport, transport aérien, transport maritime, import export, achat revente de véhicule tant en France qu'à l'étranger pouvant être nécessaires ou utiles à la réalisation de l'objet social.": "{{objet_social_ligne2}}",
    
    # SIÈGE
    "ADRESSE": "{{adresse_siege_complete}}",
    
    # DURÉE
    "99 ans": "{{duree_annees}} ans",
    
    # EXERCICE
    "31 décembre": "{{date_cloture}}",
    "31 décembre 2025": "{{premier_exercice_fin}}",
    
    # CAPITAL
    "0 000 €": "{{capital_social_formate}} €",
    "2 700 euros": "{{capital_libere_formate}} euros",
    "100 actions": "{{nombre_actions}} actions",
    "0 euros chacune": "{{valeur_nominale_formate}} euros chacune",
    
    # DATE SIGNATURE
    "Le 11 septembre 2025": "Le {{date_signature}}",
    
    # FORME JURIDIQUE
    "société par actions simplifiée unipersonnelle": "{{forme_juridique_complete}}",
    "S.A.S.U": "{{forme_juridique_sigle}}",
}

def replace_in_paragraph(paragraph):
    """Remplace le texte dans un paragraphe en préservant le formatage"""
    for run in paragraph.runs:
        for old_text, new_text in replacements.items():
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

def replace_in_tables(tables):
    """Remplace le texte dans les tableaux"""
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

# Remplacer dans tous les paragraphes
print("🔄 Remplacement des valeurs par des placeholders...")
for paragraph in doc.paragraphs:
    replace_in_paragraph(paragraph)

# Remplacer dans les tableaux
replace_in_tables(doc.tables)

# Créer le dossier templates s'il n'existe pas
os.makedirs("./templates", exist_ok=True)

# Sauvegarder le template
print(f"💾 Sauvegarde du template dans {output_file}...")
doc.save(output_file)

print("✅ Template créé avec succès !")
print(f"📁 Emplacement : {os.path.abspath(output_file)}")
print("\n🎯 Prochaine étape : Installer les dépendances npm")