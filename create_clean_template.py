from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_para(doc: Document, text: str, bold: bool = False, align: str = "left"):
    """Ajoute un paragraphe simple"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    if bold:
        run.bold = True

    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return paragraph

def create_template() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    add_para(doc, "STATUTS", bold=True, align="center")
    add_para(doc, "Société par Actions Simplifiée", align="center")
    doc.add_paragraph()

    add_para(doc, "Les soussignés :", bold=True)
    add_para(doc, "{{associe_civilite}} {{associe_prenom}} {{associe_nom}}")
    add_para(doc, "Né(e) le {{associe_date_naissance}} à {{associe_lieu_naissance}}")
    add_para(doc, "Demeurant {{associe_adresse_complete}}")
    doc.add_paragraph()
    add_para(
        doc,
        "Ont établi ainsi qu'il suit les statuts d'une société par actions simplifiée qu'ils déclarent constituer entre eux.",
    )
    doc.add_paragraph()

    add_para(doc, "TITRE I : FORME - DÉNOMINATION - SIÈGE - OBJET - DURÉE", bold=True)
    doc.add_paragraph()

    add_para(doc, "Article 1 - FORME", bold=True)
    add_para(
        doc,
        "Il est formé une {{forme_juridique_complete}} régie par les dispositions législatives et réglementaires en vigueur, et par les présents statuts.",
    )
    doc.add_paragraph()

    add_para(doc, "Article 2 - DÉNOMINATION", bold=True)
    add_para(doc, "La société a pour dénomination sociale : {{denomination}}")
    add_para(
        doc,
        "Dans tous les actes, la dénomination doit être précédée ou suivie de {{forme_juridique_complete}} ou {{forme_juridique_sigle}} et du montant du capital.",
    )
    doc.add_paragraph()

    add_para(doc, "Article 3 - SIÈGE SOCIAL", bold=True)
    add_para(doc, "Le siège social est fixé à : {{adresse_siege_complete}}")
    add_para(
        doc,
        "Il pourra être transféré par décision du président, sous réserve de ratification par l'assemblée générale.",
    )
    doc.add_paragraph()

    add_para(doc, "Article 4 - OBJET", bold=True)
    add_para(doc, "La société a pour objet :")
    add_para(doc, "{{objet_social}}")
    add_para(doc, "Et toutes opérations se rattachant directement ou indirectement à cet objet.")
    doc.add_paragraph()

    add_para(doc, "Article 5 - DURÉE", bold=True)
    add_para(doc, "La durée de la société est fixée à {{duree_societe}} années à compter de son immatriculation au RCS.")
    doc.add_paragraph()

    add_para(doc, "TITRE II : APPORTS - CAPITAL SOCIAL", bold=True)
    doc.add_paragraph()

    add_para(doc, "Article 6 - APPORTS", bold=True)
    add_para(doc, "Les associés font apport à la société de {{capital_social_formate}} euros en numéraire.")
    doc.add_paragraph()

    add_para(doc, "Article 7 - CAPITAL SOCIAL", bold=True)
    add_para(doc, "Le capital social est fixé à {{capital_social_formate}} euros.")
    add_para(
        doc,
        "Il est divisé en {{nombre_actions}} actions de {{valeur_nominale_formate}} euros chacune, libérées à hauteur de {{montant_libere_formate}} euros.",
    )
    doc.add_paragraph()

    add_para(doc, "Article 8 - MODIFICATION DU CAPITAL", bold=True)
    add_para(doc, "Le capital social peut être augmenté, réduit ou amorti dans les conditions légales.")
    doc.add_paragraph()

    add_para(doc, "Article 9 - FORME DES ACTIONS", bold=True)
    add_para(doc, "Les actions sont nominatives et donnent lieu à une inscription en compte.")
    doc.add_paragraph()

    add_para(doc, "Article 10 - TRANSMISSION DES ACTIONS", bold=True)
    add_para(doc, "La cession des actions est libre entre associés. Elle est soumise à agrément pour les tiers.")
    doc.add_paragraph()

    add_para(doc, "TITRE III : DIRECTION", bold=True)
    doc.add_paragraph()

    add_para(doc, "Article 11 - PRÉSIDENCE", bold=True)
    add_para(doc, "La société est représentée par un président.")
    add_para(doc, "Le premier président est : {{associe_civilite}} {{associe_prenom}} {{associe_nom}}")
    add_para(doc, "Le président est nommé pour une durée illimitée.")
    doc.add_paragraph()

    add_para(doc, "Article 12 - POUVOIRS DU PRÉSIDENT", bold=True)
    add_para(
        doc,
        "Le président est investi des pouvoirs les plus étendus pour agir au nom de la société dans la limite de l'objet social.",
    )
    doc.add_paragraph()

    add_para(doc, "Article 13 - RÉMUNÉRATION", bold=True)
    add_para(doc, "L'assemblée générale peut allouer au président une rémunération.")
    doc.add_paragraph()

    add_para(doc, "TITRE IV : DÉCISIONS COLLECTIVES", bold=True)
    doc.add_paragraph()

    add_para(doc, "Article 14 - ASSEMBLÉES GÉNÉRALES", bold=True)
    add_para(doc, "Les associés sont réunis en assemblée au moins une fois par an pour l'approbation des comptes.")
    doc.add_paragraph()

    add_para(doc, "Article 15 - CONVOCATION", bold=True)
    add_para(doc, "Les associés sont convoqués par le président par tous moyens 7 jours avant l'assemblée.")
    doc.add_paragraph()

    add_para(doc, "Article 16 - QUORUM ET MAJORITÉ", bold=True)
    add_para(doc, "Chaque action donne droit à une voix. Les décisions sont prises à la majorité des voix exprimées.")
    doc.add_paragraph()

    add_para(doc, "TITRE V : EXERCICE SOCIAL - COMPTES", bold=True)
    doc.add_paragraph()

    add_para(doc, "Article 17 - EXERCICE SOCIAL", bold=True)
    add_para(doc, "L'exercice social commence le 1er janvier et se termine le 31 décembre.")
    add_para(doc, "Le premier exercice se terminera le {{premier_exercice_fin}}.")
    doc.add_paragraph()

    add_para(doc, "Article 18 - COMPTES ANNUELS", bold=True)
    add_para(doc, "Le président établit les comptes annuels qui sont soumis à l'approbation de l'assemblée.")
    doc.add_paragraph()

    add_para(doc, "Article 19 - AFFECTATION DU RÉSULTAT", bold=True)
    add_para(doc, "Le bénéfice est réparti entre les associés proportionnellement au nombre d'actions.")
    doc.add_paragraph()

    add_para(doc, "TITRE VI : DISSOLUTION - LIQUIDATION", bold=True)
    doc.add_paragraph()

    add_para(doc, "Article 20 - DISSOLUTION", bold=True)
    add_para(doc, "La société prend fin par l'arrivée du terme, par décision de l'assemblée, ou pour toute cause légale.")
    doc.add_paragraph()

    add_para(doc, "Article 21 - LIQUIDATION", bold=True)
    add_para(doc, "En cas de dissolution, un ou plusieurs liquidateurs sont désignés par l'assemblée.")
    doc.add_paragraph()

    doc.add_page_break()
    add_para(doc, "Fait à {{adresse_siege_complete}}", bold=True)
    add_para(doc, "Le {{date_signature}}", bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "Signature :", bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "{{associe_prenom}} {{associe_nom}}")

    doc.save("templates/template-statuts-final.docx")

    print("✅ Template créé : templates/template-statuts-final.docx")
    print("\n📋 18 placeholders utilisés:")
    placeholders = [
        "associe_civilite",
        "associe_prenom",
        "associe_nom",
        "associe_date_naissance",
        "associe_lieu_naissance",
        "associe_adresse_complete",
        "forme_juridique_complete",
        "forme_juridique_sigle",
        "denomination",
        "adresse_siege_complete",
        "objet_social",
        "duree_societe",
        "capital_social_formate",
        "nombre_actions",
        "valeur_nominale_formate",
        "montant_libere_formate",
        "premier_exercice_fin",
        "date_signature",
    ]
    for placeholder in placeholders:
        print(f"  - {{{{{placeholder}}}}}")

if __name__ == "__main__":
    create_template()
