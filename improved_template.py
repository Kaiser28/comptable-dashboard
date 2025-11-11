# -*- coding: utf-8 -*-
from __future__ import annotations

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

SOURCE_PATH = "templates/template-statuts.docx"
OUTPUT_PATH = "templates/template-statuts-final.docx"

# Remplacements simples (ordre important pour éviter les collisions)
REPLACEMENTS = [
    (
        "Monsieur DIAOU Mamadou Né le (date) à LIEU (FRANCE) Demeurant au ({{adresse_siege_complete}})",
        "{{president_civilite}} {{president_prenom}} {{president_nom}} né(e) le {{president_date_naissance}} à {{president_lieu_naissance}} demeurant au {{president_adresse}}",
    ),
    (
        "Monsieur DIAOU Mamadou Né (date) à Ville (FRANCE)",
        "{{associe_civilite}} {{associe_prenom}} {{associe_nom}} né(e) le {{associe_date_naissance}} à {{associe_lieu_naissance}}",
    ),
    (
        "Monsieur DIAOU Mamadou    0 000 ",
        "{{associe_prenom}} {{associe_nom}}    {{associe_montant_apport_formate}} ",
    ),
    (
        "Montant des apports en numéraire  0 000 ",
        "Montant des apports en numéraire  {{total_apports_numeraires_formate}} ",
    ),
    (
        "Ces apports ont été libérés à hauteur de 2 700 euros",
        "Ces apports ont été libérés à hauteur de {{montant_libere_formate}} euros",
    ),
    (
        "Le capital social est fixé à la somme de (sommes) euros.",
        "Le capital social est fixé à la somme de {{capital_social_formate}} euros.",
    ),
    (
        "Il est divisé en 100 actions de 0 euros chacune de valeur nominale",
        "Il est divisé en {{nombre_actions_total}} actions de {{valeur_nominale_formate}} euros chacune de valeur nominale",
    ),
    (
        "31 décembre 2025",
        "{{date_premier_exercice_fin}}",
    ),
    (
        "1er janvier",
        "{{date_exercice_debut}}",
    ),
    (
        "31 décembre",
        "{{date_exercice_fin}}",
    ),
    (
        "99 ans",
        "{{duree_societe}} ans",
    ),
    (
        "Société par actions simplifiée unipersonnelle au capital de (o)euros Siège social : (adresse)",
        "{{forme_juridique_longue}} au capital de {{capital_social_formate}}  Siège social : {{adresse_siege_complete}}",
    ),
    (
        "RCS en cours dimmatriculation",
        "{{rcs_mention}}",
    ),
    (
        "Le 11 septembre 2025",
        "{{date_signature_longue}}",
    ),
    (
        "Madame DIAOU Mamadou",
        "{{mandataire_civilite}} {{mandataire_prenom}} {{mandataire_nom}}",
    ),
    (
        "SAHEL TRANSPORT",
        "{{denomination}}",
    ),
    (
        "Monsieur DIAOU Mamadou",
        "{{signataire_civilite}} {{signataire_prenom}} {{signataire_nom}}",
    ),
]

OBJECT_SOCIAL_PREFIXES = (
    "La société a pour objet",
    "Territoires d'Outre-mer",
    "- Exportation de marchandises",
    "Et plus généralement",
)


def iter_paragraphs(document: Document):
    """Iterate over all paragraphs in the document, including those in tables."""

    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def apply_replacements(paragraphs):
    for paragraph in paragraphs:
        original_text = paragraph.text

        # Gestion spécifique de l'objet social
        stripped = original_text.strip()
        if stripped.startswith(OBJECT_SOCIAL_PREFIXES[0]):
            paragraph.text = "{{objet_social}}"
            continue
        if stripped.startswith(OBJECT_SOCIAL_PREFIXES[1]) or stripped.startswith(OBJECT_SOCIAL_PREFIXES[2]) or stripped.startswith(OBJECT_SOCIAL_PREFIXES[3]):
            paragraph.text = ""
            continue

        new_text = original_text
        for search, replacement in REPLACEMENTS:
            if search in new_text:
                new_text = new_text.replace(search, replacement)

        if new_text != original_text:
            paragraph.text = new_text


def main():
    try:
        document = Document(SOURCE_PATH)
    except PackageNotFoundError as exc:
        raise SystemExit(f"Impossible d'ouvrir le document source : {exc}") from exc

    apply_replacements(iter_paragraphs(document))

    document.save(OUTPUT_PATH)
    print(f"Template enrichi sauvegardé dans {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
